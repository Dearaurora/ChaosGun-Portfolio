#!/usr/bin/env python3
"""Build the ChaosGun production playbook DOCX from its Markdown truth source.

The Markdown file is authoritative. This renderer intentionally uses a compact
reference-guide layout and normalizes the output ZIP timestamps so that the same
source and dependency versions produce byte-stable packages.
"""

from __future__ import annotations

import argparse
import math
import re
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "docs" / "workflow" / "chaosgun-demo-production-playbook.md"
DEFAULT_OUTPUT = ROOT / "docs" / "workflow" / "chaosgun-demo-production-playbook.docx"

PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_IN = 1.0
CONTENT_WIDTH_IN = PAGE_WIDTH_IN - (2 * MARGIN_IN)
CONTENT_WIDTH_DXA = 9360
CELL_MARGIN_DXA = 120

BODY_FONT = "Microsoft YaHei"
LATIN_FONT = "Aptos"
CODE_FONT = "Consolas"
ACCENT = RGBColor(0x2E, 0x74, 0xB5)
ACCENT_DARK = RGBColor(0x1F, 0x4D, 0x78)
TEXT = RGBColor(0x24, 0x2A, 0x30)
MUTED = RGBColor(0x5B, 0x66, 0x70)
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D9E1E8"


def set_run_font(run, name: str, size: float | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margin_dxa: int = CELL_MARGIN_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_dxa))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa: Sequence[int]):
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl_pr = table._tbl.tblPr

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), BODY_FONT)
    r_pr.append(r_fonts)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_TOKEN = re.compile(
    r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\((?:https?://[^)]+|#[^)]+)\))"
)


def add_inline(paragraph, text: str, *, size: float | None = None):
    """Render a deliberately small, deterministic subset of Markdown inline."""
    position = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, BODY_FONT, size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, BODY_FONT, size, True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, CODE_FONT, (size or 10.0) - 0.5)
            run.font.color.rgb = ACCENT_DARK
        else:
            link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                label, target = link_match.groups()
                if target.startswith("http"):
                    add_hyperlink(paragraph, label, target)
                else:
                    run = paragraph.add_run(label)
                    set_run_font(run, BODY_FONT, size)
                    run.font.color.rgb = ACCENT
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, BODY_FONT, size)


def create_numbering(document: Document, *, bullet: bool) -> int:
    numbering = document.part.numbering_part.element
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId"))
    ]
    num_id = max(num_ids, default=0) + 1
    desired_format = "bullet" if bullet else "decimal"
    abstract_id = None
    for abstract in numbering.findall(qn("w:abstractNum")):
        first_level = abstract.find(qn("w:lvl"))
        if first_level is None:
            continue
        num_format = first_level.find(qn("w:numFmt"))
        if num_format is not None and num_format.get(qn("w:val")) == desired_format:
            abstract_id = abstract.get(qn("w:abstractNumId"))
            break
    if abstract_id is None:
        raise RuntimeError(f"Word template has no {desired_format} numbering definition")

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    if not bullet:
        level_override = OxmlElement("w:lvlOverride")
        level_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        level_override.append(start_override)
        num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int = 0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def set_repeat_header_footer(section, title: str):
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(f"{title}  ·  v1.0")
    set_run_font(run, BODY_FONT, 8.0)
    run.font.color.rgb = MUTED
    p_pr = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "D9E1E8")
    border.append(bottom)
    p_pr.append(border)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("ChaosGun · 证据截止 2026-07-23 · ")
    set_run_font(run, BODY_FONT, 8.0)
    run.font.color.rgb = MUTED
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    field_instruction = OxmlElement("w:instrText")
    field_instruction.set(qn("xml:space"), "preserve")
    field_instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(field_instruction)
    run._r.append(field_end)


def configure_styles(document: Document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Title", 28, ACCENT_DARK, 0, 10),
        ("Subtitle", 12, MUTED, 0, 8),
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT_DARK, 14, 7),
        ("Heading 3", 12, ACCENT_DARK, 10, 5),
    ):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    caption = styles["Caption"]
    caption.font.name = BODY_FONT
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = CODE_FONT
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), CODE_FONT)
    code_style.font.size = Pt(7.5)
    code_style.font.color.rgb = RGBColor(0x26, 0x31, 0x38)
    code_style.paragraph_format.left_indent = Inches(0.12)
    code_style.paragraph_format.right_indent = Inches(0.06)
    code_style.paragraph_format.space_before = Pt(0)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = 1.0
    code_style.paragraph_format.widow_control = False

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = BODY_FONT
    callout._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = ACCENT_DARK
    callout.paragraph_format.left_indent = Inches(0.2)
    callout.paragraph_format.right_indent = Inches(0.1)
    callout.paragraph_format.space_before = Pt(7)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.2


def add_paragraph_shading(paragraph, fill: str, border_color: str | None = None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border_color:
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "14")
        left.set(qn("w:color"), border_color)
        p_bdr.append(left)
        p_pr.append(p_bdr)


def font(size: int, *, bold: bool = False):
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold else Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def centered_text(draw: ImageDraw.ImageDraw, box, text: str, text_font, fill):
    left, top, right, bottom = box
    wrapped = []
    for segment in text.split("\n"):
        wrapped.append(segment)
    line_gap = int(text_font.size * 0.25) if hasattr(text_font, "size") else 4
    heights = []
    widths = []
    for line in wrapped:
        bounds = draw.textbbox((0, 0), line, font=text_font)
        widths.append(bounds[2] - bounds[0])
        heights.append(bounds[3] - bounds[1])
    total_height = sum(heights) + line_gap * max(0, len(wrapped) - 1)
    y = top + (bottom - top - total_height) / 2
    for line, width, height in zip(wrapped, widths, heights):
        x = left + (right - left - width) / 2
        draw.text((x, y), line, font=text_font, fill=fill)
        y += height + line_gap


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, fill=(46, 116, 181), width=8):
    draw.line((start, end), fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_len = 24
    spread = math.pi / 7
    points = [
        end,
        (
            end[0] - arrow_len * math.cos(angle - spread),
            end[1] - arrow_len * math.sin(angle - spread),
        ),
        (
            end[0] - arrow_len * math.cos(angle + spread),
            end[1] - arrow_len * math.sin(angle + spread),
        ),
    ]
    draw.polygon(points, fill=fill)


def create_control_flow(path: Path):
    image = Image.new("RGB", (1800, 720), "white")
    draw = ImageDraw.Draw(image)
    title_font = font(38, bold=True)
    box_font = font(25, bold=True)
    small_font = font(20)
    draw.text((70, 36), "ChaosGun Demo 生产控制流", font=title_font, fill=(31, 77, 120))

    labels = [
        "玩家承诺\n单一风险",
        "风险原型",
        "核心循环",
        "白模与\n固定镜头",
        "Hero 切片",
        "硬化与\n最小门禁",
        "Golden / RC",
        "证据冻结\n复盘记忆",
    ]
    box_w, box_h = 185, 115
    gap = 28
    x0, y = 55, 235
    boxes = []
    for index, label in enumerate(labels):
        x = x0 + index * (box_w + gap)
        box = (x, y, x + box_w, y + box_h)
        boxes.append(box)
        fill = (234, 242, 248) if index not in (6, 7) else (222, 239, 224)
        draw.rounded_rectangle(box, radius=18, fill=fill, outline=(46, 116, 181), width=4)
        centered_text(draw, box, label, box_font, (36, 42, 48))
        if index:
            previous = boxes[index - 1]
            draw_arrow(draw, (previous[2] + 3, y + box_h / 2), (box[0] - 7, y + box_h / 2))

    # Failure loop
    source = boxes[5]
    target = boxes[4]
    loop_y = y + box_h + 125
    draw.line(
        (source[0] + box_w / 2, source[3], source[0] + box_w / 2, loop_y),
        fill=(196, 75, 69),
        width=7,
    )
    draw.line(
        (source[0] + box_w / 2, loop_y, target[0] + box_w / 2, loop_y),
        fill=(196, 75, 69),
        width=7,
    )
    draw_arrow(
        draw,
        (target[0] + box_w / 2, loop_y),
        (target[0] + box_w / 2, target[3] + 5),
        fill=(196, 75, 69),
        width=7,
    )
    draw.text((target[0] + 25, loop_y + 20), "失败指纹 ≤ 2：修复后重跑", font=small_font, fill=(150, 48, 44))
    draw.text((980, 550), "同类第 3 次：改变输入 / 实现策略 / 升级人工", font=small_font, fill=(150, 48, 44))
    image.save(path)


def create_trust_pyramid(path: Path):
    image = Image.new("RGB", (1500, 1060), "white")
    draw = ImageDraw.Draw(image)
    title_font = font(40, bold=True)
    label_font = font(27, bold=True)
    note_font = font(21)
    draw.text((70, 36), "自动化证据 → 人工裁决可信度金字塔", font=title_font, fill=(31, 77, 120))

    levels = [
        ("1  静态契约", "覆盖广 / 成本低"),
        ("2  场景·GLB·碰撞结构", "发现导入前错误"),
        ("3  Headless·AI batch", "证明可重复运行"),
        ("4  固定镜头·Golden / Rejected", "证明实际画面差异"),
        ("5  原生性能·长期运行", "证明发布环境预算"),
        ("6  真人试玩·视觉接受", "判断手感与主次"),
        ("7  发布签字·可回滚 RC", "最终责任裁决"),
    ]
    center_x = 750
    base_y = 960
    level_h = 105
    base_w = 1320
    top_w = 390
    colors = [
        (232, 239, 245),
        (218, 233, 244),
        (201, 224, 240),
        (183, 214, 235),
        (162, 203, 229),
        (125, 181, 218),
        (80, 150, 202),
    ]
    for index, (label, note) in enumerate(levels):
        y_bottom = base_y - index * level_h
        y_top = y_bottom - level_h + 4
        fraction_bottom = index / len(levels)
        fraction_top = (index + 1) / len(levels)
        width_bottom = base_w - (base_w - top_w) * fraction_bottom
        width_top = base_w - (base_w - top_w) * fraction_top
        polygon = [
            (center_x - width_top / 2, y_top),
            (center_x + width_top / 2, y_top),
            (center_x + width_bottom / 2, y_bottom),
            (center_x - width_bottom / 2, y_bottom),
        ]
        draw.polygon(polygon, fill=colors[index], outline=(255, 255, 255))
        label_box = (
            center_x - width_top / 2,
            y_top + 4,
            center_x + width_top / 2,
            y_bottom - 32,
        )
        centered_text(draw, label_box, label, label_font, (28, 55, 76))
        note_bounds = draw.textbbox((0, 0), note, font=note_font)
        draw.text(
            (center_x - (note_bounds[2] - note_bounds[0]) / 2, y_bottom - 34),
            note,
            font=note_font,
            fill=(70, 86, 98),
        )
    image.save(path)


def create_tool_radar(path: Path):
    labels = ["首次可用", "确定性", "Godot 适配", "版本控制", "协作扩展", "成本友好"]
    series = {
        "当前核心栈": ([4, 4, 5, 3, 2, 5], (85, 140, 190)),
        "近期升级栈": ([4, 5, 5, 5, 4, 4], (48, 155, 103)),
        "重型工作室栈": ([2, 5, 4, 5, 5, 1], (196, 102, 67)),
    }
    image = Image.new("RGB", (1400, 1200), "white")
    draw = ImageDraw.Draw(image, "RGBA")
    title_font = font(38, bold=True)
    label_font = font(25, bold=True)
    legend_font = font(24)
    draw.text((55, 36), "工具决策雷达（1–5，越外越优）", font=title_font, fill=(31, 77, 120))

    center = (700, 620)
    radius = 400
    count = len(labels)
    for score in range(1, 6):
        points = []
        for i in range(count):
            angle = -math.pi / 2 + (2 * math.pi * i / count)
            r = radius * score / 5
            points.append((center[0] + r * math.cos(angle), center[1] + r * math.sin(angle)))
        draw.line(points + [points[0]], fill=(175, 186, 196, 180), width=2)
    for i, label in enumerate(labels):
        angle = -math.pi / 2 + (2 * math.pi * i / count)
        end = (center[0] + radius * math.cos(angle), center[1] + radius * math.sin(angle))
        draw.line((center, end), fill=(175, 186, 196, 180), width=2)
        lx = center[0] + (radius + 75) * math.cos(angle)
        ly = center[1] + (radius + 75) * math.sin(angle)
        bounds = draw.textbbox((0, 0), label, font=label_font)
        draw.text(
            (lx - (bounds[2] - bounds[0]) / 2, ly - (bounds[3] - bounds[1]) / 2),
            label,
            font=label_font,
            fill=(36, 42, 48),
        )

    for name, (values, color) in series.items():
        points = []
        for i, score in enumerate(values):
            angle = -math.pi / 2 + (2 * math.pi * i / count)
            r = radius * score / 5
            points.append((center[0] + r * math.cos(angle), center[1] + r * math.sin(angle)))
        draw.polygon(points, fill=(*color, 45), outline=(*color, 240))
        draw.line(points + [points[0]], fill=(*color, 255), width=6)
        for point in points:
            draw.ellipse((point[0] - 8, point[1] - 8, point[0] + 8, point[1] + 8), fill=(*color, 255))

    legend_y = 1080
    legend_x = 200
    for name, (_, color) in series.items():
        draw.rounded_rectangle((legend_x, legend_y, legend_x + 34, legend_y + 24), radius=4, fill=(*color, 255))
        draw.text((legend_x + 46, legend_y - 4), name, font=legend_font, fill=(36, 42, 48))
        legend_x += 360
    image.save(path)


def add_picture(document: Document, image_path: Path, alt: str):
    with Image.open(image_path) as image:
        width, height = image.size
    max_width = CONTENT_WIDTH_IN
    max_height = 6.8
    scale = min(max_width / width, max_height / height)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width * scale))
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    # Word drawing alternative text.
    drawings = run._r.xpath(".//wp:docPr")
    if drawings:
        drawings[0].set("descr", alt)


def table_widths(rows: Sequence[Sequence[str]]) -> list[int]:
    columns = max(len(row) for row in rows)
    weights = []
    for column in range(columns):
        values = [row[column] if column < len(row) else "" for row in rows]
        max_length = max((len(re.sub(r"[`*]", "", value)) for value in values), default=1)
        weights.append(max(7, min(32, max_length)))
    total = sum(weights)
    widths = [int(CONTENT_WIDTH_DXA * weight / total) for weight in weights]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    minimum_by_columns = {
        2: 1800,
        3: 1500,
        4: 1250,
        5: 900,
        6: 750,
        7: 620,
        8: 540,
    }
    minimum = minimum_by_columns.get(columns, 500)
    for index in range(columns):
        if widths[index] >= minimum:
            continue
        deficit = minimum - widths[index]
        widths[index] = minimum
        donors = sorted(
            (candidate for candidate in range(columns) if candidate != index),
            key=lambda candidate: widths[candidate],
            reverse=True,
        )
        for donor in donors:
            available = max(0, widths[donor] - minimum)
            take = min(deficit, available)
            widths[donor] -= take
            deficit -= take
            if deficit == 0:
                break
        if deficit:
            widths[index] -= deficit
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def parse_table(lines: Sequence[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        values = [value.strip() for value in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            continue
        rows.append(values)
    return rows


def add_table(document: Document, rows: Sequence[Sequence[str]]):
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    widths = table_widths(rows)

    font_size = 8.7
    if column_count >= 7:
        font_size = 7.4
    elif column_count >= 5:
        font_size = 8.0

    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        prevent_row_split(row)
        for column_index in range(column_count):
            cell = row.cells[column_index]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline(
                paragraph,
                values[column_index] if column_index < len(values) else "",
                size=font_size,
            )
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = ACCENT_DARK
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F8FAFB")
    repeat_table_header(table.rows[0])
    set_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def extract_headings(lines: Sequence[str]) -> list[tuple[int, str]]:
    headings = []
    in_code = False
    for line in lines:
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if match:
            level = len(match.group(1))
            text = re.sub(r"[`*]", "", match.group(2)).strip()
            headings.append((level, text))
    return headings


def add_cover(document: Document, title: str):
    # Editorial rule line.
    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(60)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "24")
    top.set(qn("w:color"), "2E74B5")
    p_bdr.append(top)
    p_pr.append(p_bdr)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = kicker.add_run("CHAOSGUN · GODOT 4.6.2 · PRODUCTION SYSTEM")
    set_run_font(run, LATIN_FONT, 10, True)
    run.font.color.rgb = ACCENT
    kicker.paragraph_format.space_after = Pt(16)

    heading = document.add_paragraph(style="Title")
    add_inline(heading, title, size=28)
    heading.paragraph_format.space_after = Pt(12)

    subtitle = document.add_paragraph(style="Subtitle")
    add_inline(
        subtitle,
        "从概念、风险原型、白模、Hero 切片到 Golden / RC 的证据驱动操作手册",
        size=13,
    )

    summary = document.add_table(rows=4, cols=2)
    summary.style = "Table Grid"
    values = [
        ("冻结展示基线", "chaosgun-demo-main-v1 → 0cbe9c4"),
        ("RC 工程锚点", "0f75645"),
        ("证据审计点", "d280e7c · 130 commits · 481 task turns"),
        ("版本 / 日期", "v1.0 · 2026-07-23"),
    ]
    for row_index, (label, value) in enumerate(values):
        summary.cell(row_index, 0).text = ""
        summary.cell(row_index, 1).text = ""
        left = summary.cell(row_index, 0).paragraphs[0]
        right = summary.cell(row_index, 1).paragraphs[0]
        add_inline(left, label, size=9.5)
        add_inline(right, value, size=9.5)
        for run in left.runs:
            run.bold = True
            run.font.color.rgb = ACCENT_DARK
        set_cell_shading(summary.cell(row_index, 0), LIGHT_BLUE)
    repeat_table_header(summary.rows[0])
    set_table_geometry(summary, [2300, CONTENT_WIDTH_DXA - 2300])

    statement = document.add_paragraph(style="Callout")
    add_inline(
        statement,
        "核心原则：自动化证明结构和可重复性；玩法判断、视觉主次与最终接受由人负责。",
        size=11,
    )
    add_paragraph_shading(statement, "EEF5FA", "2E74B5")
    statement.paragraph_format.space_before = Pt(32)

    document.add_paragraph().paragraph_format.space_after = Pt(24)
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = note.add_run("维护方式")
    set_run_font(run, BODY_FONT, 9, True)
    run.font.color.rgb = ACCENT
    note.add_run("\nMarkdown 是唯一内容真源；DOCX 由仓库脚本确定性再生。")
    for run in note.runs[1:]:
        set_run_font(run, BODY_FONT, 9)
        run.font.color.rgb = MUTED

    document.add_page_break()


def add_contents(document: Document, headings: Sequence[tuple[int, str]]):
    document.add_heading("目录与使用路径", level=1)
    callout = document.add_paragraph(style="Callout")
    add_inline(
        callout,
        "第一次执行：先读第 1、3、6、8 章；需要追溯时再进入附录 A–H。",
        size=10.5,
    )
    add_paragraph_shading(callout, "EEF5FA", "2E74B5")
    bullet_num = create_numbering(document, bullet=True)
    for level, heading in headings:
        if level != 1 or heading == headings[0][1]:
            continue
        paragraph = document.add_paragraph()
        apply_numbering(paragraph, bullet_num, 0)
        add_inline(paragraph, heading, size=10.5)
        paragraph.paragraph_format.space_after = Pt(3)
    document.add_page_break()


def add_section_break_marker(document: Document):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), MID_GRAY)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def render_markdown(document: Document, source: Path, temp_dir: Path):
    lines = source.read_text(encoding="utf-8").splitlines()
    headings = extract_headings(lines)
    title = headings[0][1]
    add_cover(document, title)
    add_contents(document, headings)

    bullet_num = create_numbering(document, bullet=True)
    numbered_num = create_numbering(document, bullet=False)
    numbered_active = False
    in_code = False
    code_language = ""
    code_line_index = 0
    skip_next_code = False
    force_page_break_before_next_table = False
    first_heading_skipped = False
    i = 0

    figure_paths = {
        "control-flow": temp_dir / "control-flow.png",
        "trust-pyramid": temp_dir / "trust-pyramid.png",
        "tool-radar": temp_dir / "tool-radar.png",
    }
    create_control_flow(figure_paths["control-flow"])
    create_trust_pyramid(figure_paths["trust-pyramid"])
    create_tool_radar(figure_paths["tool-radar"])

    while i < len(lines):
        line = lines[i]

        figure_match = re.match(r"<!--\s*DOCX_FIGURE:([a-z-]+)\s*-->", line)
        if figure_match:
            key = figure_match.group(1)
            add_picture(document, figure_paths[key], key)
            skip_next_code = key in {"control-flow", "trust-pyramid"}
            force_page_break_before_next_table = key == "tool-radar"
            i += 1
            continue

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_language = line[3:].strip()
                code_line_index = 0
                if skip_next_code:
                    while i + 1 < len(lines):
                        i += 1
                        if lines[i].startswith("```"):
                            break
                    skip_next_code = False
                    in_code = False
                    code_language = ""
                i += 1
                continue
            in_code = False
            code_language = ""
            i += 1
            continue

        if in_code:
            paragraph = document.add_paragraph(style="Code Block")
            run = paragraph.add_run(line if line else " ")
            set_run_font(run, CODE_FONT, 7.5)
            add_paragraph_shading(paragraph, LIGHT_GRAY)
            if code_line_index == 0 and i + 1 < len(lines):
                paragraph.paragraph_format.keep_with_next = not lines[i + 1].startswith("```")
            code_line_index += 1
            i += 1
            continue

        if not line.strip():
            numbered_active = False
            i += 1
            continue

        if line.strip() == "---":
            add_section_break_marker(document)
            i += 1
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            if not first_heading_skipped:
                first_heading_skipped = True
                i += 1
                continue
            if level == 1:
                # Keep the compact guide dense: the first operating chapter and
                # each evidence appendix start on a fresh page, while chapters
                # 2–12 flow after their divider instead of leaving half-empty
                # pages.
                if re.match(r"^1\.", heading_text) or heading_text.startswith("附录"):
                    document.add_page_break()
            paragraph = document.add_heading(level=level)
            add_inline(paragraph, heading_text, size={1: 16, 2: 13, 3: 12}[level])
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            if force_page_break_before_next_table:
                document.add_page_break()
                force_page_break_before_next_table = False
            add_table(document, parse_table(table_lines))
            continue

        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if image_match:
            alt, target = image_match.groups()
            image_path = (source.parent / target).resolve()
            if not image_path.exists():
                raise FileNotFoundError(f"Markdown image does not exist: {image_path}")
            add_picture(document, image_path, alt)
            i += 1
            continue

        blockquote_match = re.match(r"^>\s?(.*)$", line)
        if blockquote_match:
            parts = [blockquote_match.group(1)]
            i += 1
            while i < len(lines) and lines[i].startswith(">"):
                parts.append(re.sub(r"^>\s?", "", lines[i]))
                i += 1
            paragraph = document.add_paragraph(style="Callout")
            add_inline(paragraph, "\n".join(parts), size=10)
            add_paragraph_shading(paragraph, "EEF5FA", "2E74B5")
            continue

        bullet_match = re.match(r"^(\s*)-\s+(.+)$", line)
        if bullet_match:
            paragraph = document.add_paragraph()
            apply_numbering(paragraph, bullet_num, 0)
            add_inline(paragraph, bullet_match.group(2), size=10.5)
            paragraph.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        numbered_match = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if numbered_match:
            if not numbered_active:
                numbered_num = create_numbering(document, bullet=False)
                numbered_active = True
            paragraph = document.add_paragraph()
            apply_numbering(paragraph, numbered_num, 0)
            add_inline(paragraph, numbered_match.group(2), size=10.5)
            paragraph.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Merge consecutive plain lines into one paragraph.
        paragraph_lines = [line.strip()]
        i += 1
        while i < len(lines):
            candidate = lines[i]
            if not candidate.strip():
                i += 1
                break
            if (
                candidate.startswith("#")
                or candidate.startswith("|")
                or candidate.startswith("```")
                or candidate.startswith(">")
                or candidate.startswith("![")
                or candidate.strip() == "---"
                or re.match(r"^\s*-\s+", candidate)
                or re.match(r"^\s*\d+\.\s+", candidate)
                or candidate.startswith("<!-- DOCX_FIGURE:")
            ):
                break
            paragraph_lines.append(candidate.strip())
            i += 1
        paragraph = document.add_paragraph()
        text = " ".join(paragraph_lines)
        add_inline(paragraph, text, size=11)
        if re.fullmatch(r"\*\*[^*]+\*\*", text):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        if re.match(r"^图\s*\d+", text) or re.match(r"^表\s*\d+", text):
            paragraph.style = document.styles["Caption"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def normalize_docx(source: Path, destination: Path):
    """Repack a DOCX with stable entry order and ZIP timestamps."""
    normalized = destination.with_suffix(".normalized.docx")
    with zipfile.ZipFile(source, "r") as input_zip, zipfile.ZipFile(
        normalized, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as output_zip:
        for name in sorted(input_zip.namelist()):
            original = input_zip.getinfo(name)
            data = input_zip.read(name)
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            # Preserve Office's package attributes. Replacing create_system or
            # external_attr can leave a DOCX readable while making Word's PDF
            # exporter stall indefinitely.
            info.compress_type = original.compress_type
            info.external_attr = original.external_attr
            info.internal_attr = original.internal_attr
            info.create_system = original.create_system
            info.create_version = original.create_version
            info.extract_version = original.extract_version
            info.flag_bits = original.flag_bits
            info.volume = original.volume
            output_zip.writestr(info, data)
    normalized.replace(destination)


def build(source: Path, output: Path):
    if not source.exists():
        raise FileNotFoundError(source)
    output.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    configure_styles(document)
    set_repeat_header_footer(section, "ChaosGun Demo 生产手册")

    fixed_time = datetime(2026, 7, 23, 0, 0, 0, tzinfo=timezone.utc)
    document.core_properties.title = "ChaosGun Godot Demo 可复制稳定工作流手册"
    document.core_properties.subject = "Evidence-driven demo production playbook"
    document.core_properties.author = "ChaosGun Project"
    document.core_properties.keywords = "Godot, Demo, Workflow, Golden, RC, Evidence"
    document.core_properties.comments = "Generated from Markdown; do not edit as a second truth source."
    document.core_properties.created = fixed_time
    document.core_properties.modified = fixed_time
    document.core_properties.revision = 1

    with tempfile.TemporaryDirectory(prefix="chaosgun-playbook-") as temp:
        temp_dir = Path(temp)
        render_markdown(document, source, temp_dir)
        raw_output = temp_dir / "raw.docx"
        document.save(raw_output)
        normalize_docx(raw_output, output)


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build(args.source.resolve(), args.output.resolve())
    print(f"Built {args.output.resolve()}")


if __name__ == "__main__":
    main()
